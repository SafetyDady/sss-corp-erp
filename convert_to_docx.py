"""
Convert SYSTEM_OVERVIEW_V3.md → Word (.docx)
Handles: Headings, Tables, Code blocks, Bold, Bullet points, Blockquotes
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_docx(md_path, docx_path):
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Heading styles ---
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0, 51, 102)
        if i == 1:
            hs.font.size = Pt(20)
        elif i == 2:
            hs.font.size = Pt(16)
        elif i == 3:
            hs.font.size = Pt(13)
        else:
            hs.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line
            run = p.add_run('─' * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(180, 180, 180)
            i += 1
            continue

        # --- Headings ---
        if line.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.*)', line)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                doc.add_heading(text, level=level)
                i += 1
                continue

        # --- Mermaid / Code blocks ---
        if line.strip().startswith('```'):
            lang = line.strip().replace('```', '').strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```

            if lang == 'mermaid':
                # Add a note that this is a diagram
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run('📊 [Mermaid Diagram — ดูใน VS Code: Ctrl+Shift+V]')
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                # Still show the code for reference
                for cl in code_lines:
                    cp = doc.add_paragraph(cl)
                    cp.style = doc.styles['Normal']
                    for run in cp.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(80, 80, 80)
                    cp.paragraph_format.space_before = Pt(0)
                    cp.paragraph_format.space_after = Pt(0)
                    cp.paragraph_format.left_indent = Cm(1)
            else:
                # ASCII art / code block — preserve as-is
                for cl in code_lines:
                    cp = doc.add_paragraph(cl)
                    for run in cp.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    cp.paragraph_format.space_before = Pt(0)
                    cp.paragraph_format.space_after = Pt(0)
                    cp.paragraph_format.left_indent = Cm(0.5)
            continue

        # --- Tables ---
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Parse table
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                # Parse header
                headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                # Parse rows (skip separator line)
                rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split('|') if c.strip()]
                    if cells:
                        rows.append(cells)

                # Create Word table
                if headers:
                    num_cols = len(headers)
                    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT

                    # Header row
                    hdr = table.rows[0]
                    for j, h in enumerate(headers):
                        if j < num_cols:
                            cell = hdr.cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(clean_md(h))
                            run.bold = True
                            run.font.size = Pt(10)
                            run.font.name = 'Calibri'
                            # Header bg color
                            shading = cell._element.get_or_add_tcPr()
                            shade_elem = shading.makeelement(qn('w:shd'), {
                                qn('w:val'): 'clear',
                                qn('w:color'): 'auto',
                                qn('w:fill'): '003366'
                            })
                            shading.append(shade_elem)
                            run.font.color.rgb = RGBColor(255, 255, 255)

                    # Data rows
                    for ri, row_data in enumerate(rows):
                        row = table.rows[ri + 1]
                        for j, val in enumerate(row_data):
                            if j < num_cols:
                                cell = row.cells[j]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                text = clean_md(val)
                                run = p.add_run(text)
                                run.font.size = Pt(10)
                                run.font.name = 'Calibri'
                                # Alternating row color
                                if ri % 2 == 0:
                                    shading = cell._element.get_or_add_tcPr()
                                    shade_elem = shading.makeelement(qn('w:shd'), {
                                        qn('w:val'): 'clear',
                                        qn('w:color'): 'auto',
                                        qn('w:fill'): 'F0F4F8'
                                    })
                                    shading.append(shade_elem)

                    # Add small space after table
                    doc.add_paragraph()
            continue

        # --- Blockquote (> ...) → highlight box ---
        if line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            # Collect multi-line quotes
            while i + 1 < len(lines) and lines[i + 1].startswith('>'):
                i += 1
                quote_text += '\n' + lines[i].lstrip('>').strip()

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

            # Add colored marker
            marker = p.add_run('💬 ')
            marker.font.size = Pt(11)

            run = p.add_run(clean_md(quote_text))
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 100, 180)
            run.bold = True
            i += 1
            continue

        # --- Bullet points ---
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # --- Bold line (standalone **text**) ---
        # --- Normal paragraph ---
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1

    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")


def clean_md(text):
    """Remove markdown formatting for plain text"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text


def add_formatted_text(paragraph, text):
    """Add text with bold/italic/code formatting"""
    # Split by **bold**, *italic*, `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 50, 50)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'


if __name__ == '__main__':
    create_docx(
        'C:/web_project/sss-corp/sss-corp-erp/SYSTEM_OVERVIEW_V3.md',
        'C:/web_project/sss-corp/sss-corp-erp/SYSTEM_OVERVIEW_V3.docx'
    )
